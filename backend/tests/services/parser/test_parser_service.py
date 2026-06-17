"""ParserService 主入口 + 各子解析器单元测试（任务 13）。

覆盖：
- PDF 文本层路径（密度 ≥ 阈值）
- PDF 密度阈值回退 OCR
- PDF 完全无文本层 → OCR
- PDF 损坏 → failed
- DOCX 正文 + 表格
- DOCX .doc 老格式 / 损坏 → failed
- 图片 → OCR
- 图片损坏 → failed
- 文本 < 50 字符 → low_text
- 不支持的 mime → failed

fixture 生成：
- PDF 用 ``fpdf2``（轻量、无系统依赖），失败回退到 ``pypdf`` 写少量文本
- DOCX 用 ``python-docx`` 直接写
- 图片用 ``Pillow``
"""
from __future__ import annotations

import io
from typing import Any

import pytest
from docx import Document
from PIL import Image, ImageDraw, ImageFont

from app.services.parser import ParserService, ParsedResult
from app.services.parser.docx_parser import DOCXParseError, parse_docx
from app.services.parser.ocr import OCRAdapter, StubOCRAdapter
from app.services.parser.pdf_parser import PDFParseError, parse_pdf


# ============================================================================
# Fixture 工厂
# ============================================================================


def _make_pdf(*, text_per_page: list[str] | None = None) -> bytes:
    """用 fpdf2 生成包含文本层的 PDF（每页一段文字）。"""
    try:
        from fpdf import FPDF
    except ImportError as exc:
        pytest.skip(f"fpdf2 not installed: {exc}")

    text_per_page = text_per_page or ["Hello World " * 20]
    pdf = FPDF()
    # 加载一个基础字体（fpdf2 自带 helvetica）
    pdf.set_font("helvetica", size=12)
    for text in text_per_page:
        pdf.add_page()
        # 多行写入，让密度足够
        for i in range(0, min(len(text), 800), 80):
            pdf.cell(0, 10, text[i : i + 80])
            pdf.ln(5)
    out = pdf.output()
    # fpdf2 返回 bytearray；转 bytes
    return bytes(out)


def _make_scanned_pdf(num_pages: int = 1) -> bytes:
    """生成"扫描版"PDF：把白色图片转 PDF（无文本层）。

    PIL 支持把多张图合并为 PDF。
    """
    images = []
    for _ in range(num_pages):
        img = Image.new("RGB", (200, 300), color="white")
        images.append(img)
    buf = io.BytesIO()
    images[0].save(buf, format="PDF", save_all=True, append_images=images[1:])
    return buf.getvalue()


def _make_docx(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    """生成一个简单的 .docx。"""
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    if table_rows:
        t = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for i, row in enumerate(table_rows):
            for j, cell in enumerate(row):
                t.cell(i, j).text = cell
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_png(text_in_image: str = "name: Zhang San phone: 13800138000") -> bytes:
    """生成一张 PNG（白色背景 + 文字）。"""
    img = Image.new("RGB", (400, 100), color="white")
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.load_default()
    except Exception:  # noqa: BLE001
        font = None
    draw.text((10, 40), text_in_image, fill="black", font=font)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


class _FakeOCR(OCRAdapter):
    """可注入的 fake OCR（测试可控）。"""

    def __init__(self, *, returns: str = "OCR_FAKE_TEXT " * 20) -> None:
        self._returns = returns
        self.calls: list[bytes] = []

    @property
    def backend_name(self) -> str:
        return "fake"

    async def extract(self, image_bytes: bytes, *, langs: tuple[str, ...]) -> str:
        self.calls.append(image_bytes)
        return self._returns


# ============================================================================
# ParserService 主入口
# ============================================================================


class TestParserServiceDispatch:
    """测试 ParserService 按 mime 路由 + 状态分类。"""

    async def test_pdf_text_layer_returns_success(self) -> None:
        pdf = _make_pdf(text_per_page=["Name Zhang San Phone 138 Email test@a.com " * 5])
        service = ParserService(ocr=StubOCRAdapter())
        result = await service.parse(pdf, mime="application/pdf")

        assert result.status == "success"
        assert "Name" in result.text

    async def test_pdf_low_text_marks_low_text_status(self) -> None:
        # 每页 < PARSE_MIN_TEXT_LENGTH 字符 → low_text
        # 同时密度也低（< 100），会走 OCR；这里用 fake OCR 也返回短文本
        fake_short = _FakeOCR(returns="short")  # 5 chars
        pdf = _make_scanned_pdf(num_pages=1)  # 无文本层 → 必走 OCR
        service = ParserService(ocr=fake_short)
        result = await service.parse(pdf, mime="application/pdf")

        assert result.status == "low_text"
        assert "short" in result.text

    async def test_image_routes_to_ocr(self) -> None:
        png = _make_png()
        fake = _FakeOCR(returns="OCR result with enough text length to be > 50 characters")
        service = ParserService(ocr=fake)
        result = await service.parse(png, mime="image/png")

        assert result.status == "success"
        assert len(fake.calls) == 1
        assert "OCR result" in result.text

    async def test_docx_extracts_paragraphs_and_tables(self) -> None:
        # 凑够 > 50 字符的正文（避免被分类为 low_text）
        long_para = "候选人姓名：张三。" + "工作经历详细描述。" * 5
        docx = _make_docx(
            paragraphs=[long_para],
            table_rows=[["技能", "经验"], ["Python", "5年"]],
        )
        service = ParserService(ocr=StubOCRAdapter())
        result = await service.parse(
            docx,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        assert result.status == "success"
        assert "张三" in result.text
        assert "Python" in result.text
        assert "5年" in result.text

    async def test_legacy_doc_format_marks_failed(self) -> None:
        service = ParserService(ocr=StubOCRAdapter())
        result = await service.parse(b"OLE2 fake", mime="application/msword")

        assert result.status == "failed"
        assert "doc" in (result.error or "").lower()

    async def test_unsupported_mime_marks_failed(self) -> None:
        service = ParserService(ocr=StubOCRAdapter())
        result = await service.parse(b"...", mime="text/csv")

        assert result.status == "failed"
        assert "unsupported" in (result.error or "").lower()

    async def test_empty_content_marks_failed(self) -> None:
        service = ParserService(ocr=StubOCRAdapter())
        result = await service.parse(b"", mime="application/pdf")

        assert result.status == "failed"
        assert result.error is not None

    async def test_corrupt_pdf_marks_failed(self) -> None:
        service = ParserService(ocr=StubOCRAdapter())
        result = await service.parse(b"not a real pdf", mime="application/pdf")

        assert result.status == "failed"
        assert result.error is not None

    async def test_corrupt_image_marks_failed(self) -> None:
        service = ParserService(ocr=StubOCRAdapter())
        result = await service.parse(b"not a real image", mime="image/png")

        assert result.status == "failed"
        assert result.error is not None


# ============================================================================
# PDF 解析器：密度阈值回退 OCR
# ============================================================================


class TestPDFDensityFallback:
    """字符密度阈值回退 OCR 的关键测试。"""

    async def test_high_density_no_ocr_invoked(self) -> None:
        # 每页 > 100 字符
        pdf = _make_pdf(text_per_page=["Hello World Content " * 30])
        fake = _FakeOCR()
        text = await parse_pdf(pdf, ocr=fake, density_threshold=100)

        assert len(fake.calls) == 0  # 文本层够，不调 OCR
        # 文本可能被 pdfplumber 加换行，去掉空白再断言
        assert "Hello" in text
        assert "World" in text
        assert "Content" in text

    async def test_low_density_triggers_ocr_fallback(self) -> None:
        """每页 < 阈值 → 触发 OCR 回退（所有页都 rasterize + OCR）。"""
        pdf = _make_pdf(text_per_page=["hi"])  # 2 字符
        fake = _FakeOCR(returns="OCR RESULT " * 10)
        text = await parse_pdf(
            pdf,
            ocr=fake,
            density_threshold=100,
            rasterizer=lambda content, idx: b"fake_png_bytes",
        )

        assert len(fake.calls) == 1  # 1 页 → 1 次 OCR
        assert "OCR RESULT" in text

    async def test_empty_pdf_raises_pdf_parse_error(self) -> None:
        with pytest.raises(PDFParseError):
            await parse_pdf(b"", ocr=StubOCRAdapter())


# ============================================================================
# DOCX 解析器
# ============================================================================


class TestDOCXParser:
    def test_extracts_paragraphs(self) -> None:
        docx = _make_docx(["Hello", "World"])
        text = parse_docx(docx)
        assert "Hello" in text
        assert "World" in text

    def test_extracts_tables(self) -> None:
        docx = _make_docx([], table_rows=[["A", "B"], ["1", "2"]])
        text = parse_docx(docx)
        assert "A\tB" in text
        assert "1\t2" in text

    def test_empty_content_raises(self) -> None:
        with pytest.raises(DOCXParseError):
            parse_docx(b"")

    def test_invalid_docx_raises(self) -> None:
        with pytest.raises(DOCXParseError):
            parse_docx(b"not a real docx")


# ============================================================================
# OCR adapter 抽象
# ============================================================================


class TestStubOCRAdapter:
    async def test_stub_returns_empty_string_with_warning(self, caplog: pytest.LogCaptureFixture) -> None:
        adapter = StubOCRAdapter(reason="test")
        result = await adapter.extract(b"image", langs=("ch", "en"))
        assert result == ""
        assert adapter.backend_name == "stub"
