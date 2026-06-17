"""DOCX 解析器（任务 13）。

用 ``python-docx`` 提取：
- 段落正文（``document.paragraphs``）
- 表格文本（``document.tables`` 每个单元格）

约束：流式累积文本；不超过 ``MAX_DOCX_TEXT_BYTES``（防 zip bomb）。

注意：``.doc`` 老格式（OLE2）不支持，需要 LibreOffice 转换 → 直接报错让上游用 OCR 路径。
"""
from __future__ import annotations

import io

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from app.core.logging import get_logger

logger = get_logger(__name__)


MAX_DOCX_TEXT_BYTES: int = 2 * 1024 * 1024  # 2MB；超限截断
"""防御性上限：超大的 docx 文本（如几百万字）也截断，避免撑爆 LLM token。"""


class DOCXParseError(Exception):
    """docx 解析错误（损坏 / 非法格式 / .doc 老格式）。"""


def parse_docx(content: bytes) -> str:
    """提取 docx 正文 + 表格，返回拼接文本。

    Args:
        content: .docx 字节流

    Returns:
        全文（段落间 ``\\n\\n``，表格内单元格用 ``\\t`` 分隔）

    Raises:
        DOCXParseError: 文件不是合法 docx（含 .doc 老格式场景）
    """
    if not content:
        raise DOCXParseError("empty docx content")

    try:
        document = Document(io.BytesIO(content))
    except PackageNotFoundError as exc:
        # 常见原因：用户上传 .doc 老格式（OLE2 复合文档）
        raise DOCXParseError(
            "not a valid .docx package (likely legacy .doc format)"
        ) from exc
    except Exception as exc:  # noqa: BLE001
        raise DOCXParseError(f"python-docx failed: {exc}") from exc

    chunks: list[str] = []
    total_bytes = 0

    for para in document.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        chunk = text + "\n"
        total_bytes += len(chunk.encode("utf-8"))
        if total_bytes > MAX_DOCX_TEXT_BYTES:
            logger.warning(
                "docx_text_truncated",
                limit_bytes=MAX_DOCX_TEXT_BYTES,
            )
            chunks.append("[truncated]")
            break
        chunks.append(chunk)

    # 表格：每行单元格用 \t 分隔
    if total_bytes <= MAX_DOCX_TEXT_BYTES:
        for table in document.tables:
            for row in table.rows:
                cells = [
                    (c.text or "").strip().replace("\n", " ") for c in row.cells
                ]
                line = "\t".join(cells) + "\n"
                total_bytes += len(line.encode("utf-8"))
                if total_bytes > MAX_DOCX_TEXT_BYTES:
                    chunks.append("[truncated]")
                    logger.warning(
                        "docx_text_truncated_in_table",
                        limit_bytes=MAX_DOCX_TEXT_BYTES,
                    )
                    break
                chunks.append(line)
            else:
                continue
            break

    return "".join(chunks).strip()


__all__ = ["parse_docx", "DOCXParseError", "MAX_DOCX_TEXT_BYTES"]
